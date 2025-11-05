from flask import Flask, render_template, request, jsonify, send_file, redirect
import os, re, json, secrets
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
import stripe

# ------------------------------------------------------------------
# 1️⃣ Parse morrowland 243.docx
# ------------------------------------------------------------------
def load_detailed_archetypes_docx(file_path: str):
    if not os.path.exists(file_path):
        print(f"[ERROR] File not found: {file_path}")
        return {}, {}
    doc = Document(file_path)
    raw_lines = [p.text for p in doc.paragraphs]
    lines = [line.strip() for line in raw_lines]

    header_re = re.compile(
        r"(?i)openness\s*[:\-–—]?\s*(low|medium|high).*?"
        r"conscientiousness\s*[:\-–—]?\s*(low|medium|high).*?"
        r"extraversion\s*[:\-–—]?\s*(low|medium|high).*?"
        r"agreeableness\s*[:\-–—]?\s*(low|medium|high).*?"
        r"neuroticism\s*[:\-–—]?\s*(low|medium|high)"
    )
    archetype_re = re.compile(r"(?i)^archetype\s*[:\-–—]?\s*(.+?)\s*$")

    by_code, by_name = {}, {}
    current_code, current_name, buffer = None, None, []

    def flush():
        nonlocal current_code, current_name, buffer
        if current_code and buffer:
            text = "\n".join(buffer).strip()
            by_code[current_code] = text
            if current_name:
                by_name[current_name] = text
        buffer = []

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        m_header = header_re.search(line)
        if m_header:
            flush()
            O, C, E, A, N_ = [x.capitalize() for x in m_header.groups()]
            current_code = f"{O}-{C}-{E}-{A}-{N_}"
            current_name = None
            for j in range(1, 4):
                if i + j >= n:
                    break
                next_line = lines[i + j].strip()
                m_name = archetype_re.match(next_line)
                if m_name:
                    current_name = m_name.group(1).strip()
                    i = i + j
                    break
            if not current_name:
                current_name = f"Unknown_{i}"
        else:
            if current_code:
                buffer.append(raw_lines[i])
        i += 1
    flush()
    print(f"[✅ SUCCESS] Loaded {len(by_code)} archetypes from {os.path.basename(file_path)}")
    return by_code, by_name


# ------------------------------------------------------------------
# 2️⃣ Setup Flask + Stripe
# ------------------------------------------------------------------
load_dotenv()
app = Flask(__name__)
stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
DOMAIN = os.getenv("DOMAIN", "http://localhost:5000")


# ------------------------------------------------------------------
# 3️⃣ Load morrowland 243.docx
# ------------------------------------------------------------------
base_dir = os.path.dirname(os.path.abspath(__file__))
file_path = os.path.join(base_dir, "morrowland 243.docx")
DETAILED_BY_CODE, DETAILED_BY_NAME = load_detailed_archetypes_docx(file_path)


# ------------------------------------------------------------------
# 4️⃣ Archetypes + free code storage
# ------------------------------------------------------------------
def load_archetypes():
    for file in ["archetypes_full.json", "archetypes.json"]:
        if os.path.exists(file):
            with open(file, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and data:
                    print(f"[INFO] Loaded {len(data)} archetypes from {file}")
                    return data
    print("[⚠️ Using fallback minimal archetypes]")
    return {"Low-Low-Low-Low-Low": "Aquashine"}


ARCHETYPES = load_archetypes()
FREE_CODES_FILE = os.path.join(base_dir, "free_codes.json")


# ------------------------------------------------------------------
# 5️⃣ One-time free access codes
# ------------------------------------------------------------------
def load_free_codes():
    if os.path.exists(FREE_CODES_FILE):
        with open(FREE_CODES_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_free_codes(codes):
    with open(FREE_CODES_FILE, "w", encoding="utf-8") as f:
        json.dump(codes, f, indent=2)


def generate_free_code():
    code = secrets.token_hex(4).upper()  # 8-char hex code
    codes = load_free_codes()
    codes[code] = {"used": False}
    save_free_codes(codes)
    print(f"[🎁 NEW FREE CODE GENERATED]: {code}")
    return code


def verify_free_code(code: str) -> bool:
    """Return True and mark used if code is valid and unused."""
    codes = load_free_codes()
    if code in codes and not codes[code]["used"]:
        codes[code]["used"] = True
        save_free_codes(codes)
        print(f"[✅ FREE CODE USED]: {code}")
        return True
    return False


# ------------------------------------------------------------------
# 6️⃣ Global social links
# ------------------------------------------------------------------
@app.context_processor
def inject_socials():
    return dict(
        tiktok_url="https://www.tiktok.com/@neptunee7777",
        instagram_url="https://www.instagram.com/kendallm16"
    )


# ------------------------------------------------------------------
# 7️⃣ Routes
# ------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/quiz")
def quiz():
    return render_template("quiz.html")


@app.route("/generate-free-code")
def make_free_code():
    """You (the owner) can hit this URL to mint a new one-time code."""
    code = generate_free_code()
    return jsonify({"new_code": code})


@app.route("/verify-free-code", methods=["POST"])
def api_verify_free_code():
    """
    Called from the quiz page via fetch().
    Expects JSON: { "code": "ABC123" }
    Returns JSON: { "valid": true/false, "error": "...optional..." }
    """
    data = request.get_json() or {}
    raw = (data.get("code") or "").strip()
    if not raw:
        return jsonify({"valid": False, "error": "No code provided."}), 400

    code = raw.upper()  # make it case-insensitive for the user

    if verify_free_code(code):
        return jsonify({"valid": True})
    else:
        return jsonify({"valid": False, "error": "Invalid or already used."}), 400


@app.route("/report")
def report():
    """
    Older flow: /report?code=...&free=YOURCODE
    Still works, but the quiz now mostly uses /verify-free-code instead.
    """
    code = request.args.get("code", "")
    free_key = (request.args.get("free", "") or "").strip().upper()
    if free_key and verify_free_code(free_key):
        return redirect(f"/api/render-report?code={code}&paid=true")
    return redirect(f"/create-checkout-session?code={code}")


@app.route("/create-checkout-session")
def create_checkout_session():
    code = request.args.get("code", "Medium-Medium-Medium-Medium-Medium")
    try:
        session = stripe.checkout.Session.create(
            mode="payment",
            payment_method_types=["card"],
            line_items=[{
                "price_data": {
                    "currency": "usd",
                    "product_data": {"name": "Big 5 Detailed Archetype Report"},
                    "unit_amount": 99,  # $0.99
                },
                "quantity": 1
            }],
            success_url=f"{DOMAIN}/api/render-report?code={code}&paid=true",
            cancel_url=f"{DOMAIN}/"
        )
        return redirect(session.url)
    except Exception as e:
        print("Stripe error:", e)
        return f"Stripe session creation failed: {e}", 500


@app.route("/api/render-report")
def api_render_report():
    code = request.args.get("code", "")
    paid = request.args.get("paid", "").lower() == "true"

    detailed_text = DETAILED_BY_CODE.get(code)
    archetype_name = None

    if not detailed_text and code in ARCHETYPES:
        archetype_name = ARCHETYPES[code]
        detailed_text = DETAILED_BY_NAME.get(archetype_name)

    if not detailed_text:
        detailed_text = "Detailed report not found."
    if not archetype_name and code in ARCHETYPES:
        archetype_name = ARCHETYPES[code]

    sections = {"Detailed Report": detailed_text} if paid else {
        "Summary": "Preview only. Purchase or use a free code to unlock the full report."
    }
    subtype = "N/A" if paid else "Locked"

    return render_template(
        "detailed_report.html",
        archetype=archetype_name or "Unknown",
        traits=code,
        subtype=subtype,
        sections=sections,
        quote="“Depth rewards patience.”"
    )


@app.route("/api/download-report")
def download_report():
    code = request.args.get("code", "")
    name = ARCHETYPES.get(code, "Unknown")
    detailed_text = DETAILED_BY_CODE.get(code) or DETAILED_BY_NAME.get(name)

    doc = Document()
    doc.add_heading(name, level=1)
    doc.add_paragraph(detailed_text or "Detailed text not found.")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name=f"{name.replace(' ', '_')}_Detailed_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ------------------------------------------------------------------
# 🏁 Run Flask
# ------------------------------------------------------------------
if __name__ == "__main__":
    app.run(debug=True)